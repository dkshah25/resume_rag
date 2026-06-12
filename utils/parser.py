import fitz  # PyMuPDF
import docx
import io
import logging
from typing import Tuple

logger = logging.getLogger("hirelens.parser")

def extract_text_from_pdf(file_bytes: bytes) -> Tuple[str, int]:
    """
    Extracts text and page count from a PDF document in bytes.
    """
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        pages = len(doc)
        for page in doc:
            text += page.get_text() + "\n"
        
        # Clean extracted text a bit
        cleaned_text = "\n".join([line.strip() for line in text.split("\n") if line.strip()])
        return cleaned_text, pages
    except Exception as e:
        logger.error(f"Failed to parse PDF file: {e}")
        raise ValueError(f"Corrupted or invalid PDF format: {str(e)}")

def extract_text_from_docx(file_bytes: bytes) -> Tuple[str, int]:
    """
    Extracts text and page count estimate from a DOCX document in bytes.
    """
    try:
        file_stream = io.BytesIO(file_bytes)
        doc = docx.Document(file_stream)
        text_elements = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_elements.append(para.text.strip())
                
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_elements.append(" | ".join(row_text))
                    
        text = "\n".join(text_elements)
        
        # Estimate page count (standard page is ~3000 characters)
        pages = max(1, len(text) // 3000)
        return text, pages
    except Exception as e:
        logger.error(f"Failed to parse DOCX file: {e}")
        raise ValueError(f"Corrupted or invalid DOCX format: {str(e)}")

def extract_content(file_bytes: bytes, filename: str) -> Tuple[str, int]:
    """
    Directs the extraction to the appropriate parser based on file extension.
    """
    if not file_bytes:
        raise ValueError("The uploaded file is empty.")
        
    lower_filename = filename.lower()
    if lower_filename.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif lower_filename.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file format: {filename}. Only PDF and DOCX are supported.")
